import docx
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def ajustar_margin(documento,margin=1.25):
    """Mudas as margens do documento"""
    sections = documento.sections
    for section in sections:
        section.top_margin = Cm(margin)
        section.bottom_margin = Cm(margin)
        section.left_margin = Cm(margin)
        section.right_margin = Cm(margin)
    return documento


def inserir_imagens(arquivos, config):

    ##-----------Tratamento de pasta e arquivos---------##
    local_arquivo = str(arquivos[0])
    arquivo_imagem = local_arquivo.split('/')[-1]
    nome_arquivo_saida = local_arquivo.split('/')[-2]
    local_imagens = local_arquivo.replace(arquivo_imagem,"")

    ##_-------------------------------------------------##

    doc = docx.Document()
    table = doc.add_table(rows=1, cols=config["grade"])
    linha = 0
    celulas = 0
    for imagens in arquivos:
        row_cells = table.rows[linha].cells
        p = row_cells[celulas].add_paragraph()

        # ----- Centralizar Imagem -------#
        if config["centralizar"]:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # --------------------------------#

        r = p.add_run()
        r.add_picture(imagens,width=Cm(config["largura_imagem"]), height=Cm(config["altura_imagem"]))
        if config["legenda"]:
            p = row_cells[celulas].add_paragraph(f"{nome_arquivo_saida}")
        celulas += 1
        if celulas % config["grade"] == 0:
            row_cells = table.add_row().cells
            celulas = 0
            linha += 1

    # --------- Mudar margens do documento ------- #
    if config['margin']:
        doc = ajustar_margin(doc)
    # ------------------------------------ ------- #

    doc.save(local_imagens+nome_arquivo_saida+'.docx')


if __name__ == "__main__":
    pass

