import docx
from gtk_dialog import open_file_dialog
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


# def inserir(arquivos):
#     local_arquivo = str(arquivos[0])
#     arquivo = local_arquivo.split('/')[-1]
    
#     local = local_arquivo.replace(arquivo,"")


#     doc = docx.Document()
#     p = doc.add_paragraph()
#     r = p.add_run()
#     r.add_picture(local_arquivo,width=Inches(4.0), height=Inches(4.0))
#     last_paragraph = doc.paragraphs[-1]
#     last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     doc.save(local+'imgens-da-familia.docx')
       


# def inserir_varios(arquivos,altura_img=3.75, largura_img=2.5):
#     doc = docx.Document()
#     local_arquivo = str(arquivos[0])
#     arquivo = local_arquivo.split('/')[-1]
#     local = local_arquivo.replace(arquivo,"")
#     tables = doc.tables
#     linha = 1
#     table = doc.add_table(rows=0, cols=2)
#     row_cells = table.add_row().cells
#     for imagens in arquivos:
#         if linha == 1:
#             fisrt_row_cells = table.rows[0].cells
#             p = fisrt_row_cells[0].add_paragraph()
#             r = p.add_run()
#             r.add_picture(imagens,width=Inches(largura_img), height=Inches(altura_img))
#         elif linha == 2:
#             fisrt_row_cells = table.rows[0].cells
#             p = fisrt_row_cells[1].add_paragraph()
#             r = p.add_run()
#             r.add_picture(imagens,width=Inches(largura_img), height=Inches(altura_img))
#         elif linha%2 !=0:
#             row_cells = table.add_row().cells
#             p = row_cells[0].add_paragraph()
#             r = p.add_run()
#             r.add_picture(imagens,width=Inches(largura_img), height=Inches(altura_img))
#         else:
#             p = row_cells[1].add_paragraph()
#             r = p.add_run()
#             r.add_picture(imagens,width=Inches(largura_img), height=Inches(altura_img))
#         linha+=1
#     doc.save(local+'imgens-da-familia.docx')




def inserir_varios(arquivos, config):
    
    ##-----------Tratamento de pasta e arquivos---------##
    local_arquivo = str(arquivos[0])
    arquivo_imagem = local_arquivo.split('/')[-1]
    nome_arquivo_saida = local_arquivo.split('/')[-2]
    local_imagens = local_arquivo.replace(arquivo_imagem,"")
    ##_-------------------------------------------------##
    
    doc = docx.Document()
    table = doc.add_table(rows=1, cols=config["grade"])
    linha = 0
    celulas = 0
    for imagens in arquivos:
        row_cells = table.rows[linha].cells
        p = row_cells[celulas].add_paragraph()
        
        # ----- Centralizar Imagem -------#
        if config["centralizar"]:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # --------------------------------#
        
        r = p.add_run()
        r.add_picture(imagens,width=Cm(config["largura_imagem"]), height=Cm(config["altura_imagem"]))
        
        celulas += 1
        if celulas % config["grade"] == 0:
            row_cells = table.add_row().cells
            celulas = 0
            linha += 1
    
    # --------- Mudar margens do documento ------- #
    if config['margin']:
        doc = ajustar_margin(doc)
    # ------------------------------------ ------- #

    doc.save(local_imagens+nome_arquivo_saida+'.docx')


def ajustar_margin(documento,margin=1.27):
    """Mudas as margens do documento"""
    sections = documento.sections
    for section in sections:
        section.top_margin = Cm(margin)
        section.bottom_margin = Cm(margin)
        section.left_margin = Cm(margin)
        section.right_margin = Cm(margin)
    return documento


configuracao = {1:{"grade":4,"largura_imagem":4,"altura_imagem":4,"centralizar":True,"margin":True},
2:{"grade":3,"largura_imagem":5,"altura_imagem":7,"centralizar":True,"margin":False},
3:{"grade":4,"largura_imagem":4.5,"altura_imagem":6,"centralizar":True,"margin":True},
4:{"grade":3,"largura_imagem":6,"altura_imagem":6,"centralizar":True,"margin":True},
5:{"grade":3,"largura_imagem":5,"altura_imagem":6.5,"centralizar":True,"margin":True}}


arquivos = open_file_dialog()


inserir_varios(arquivos,configuracao[3])
# if len(arquivos) > 2:
#     inserir_varios(arquivos,configuracao[2])
# else:
#     inserir_varios(arquivos,configuracao[len(arquivos)])
    

